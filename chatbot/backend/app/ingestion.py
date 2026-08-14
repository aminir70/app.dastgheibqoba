import io
import logging
import re

import fitz  # PyMuPDF
import pytesseract
import tiktoken
from PIL import Image
from docx import Document as DocxDocument

from .db import SessionLocal
from .models import Chunk, Document
from .openai_client import embed_texts
from .settings_store import get_setting
from .text_norm import normalize_fa

log = logging.getLogger("ingestion")

# A page with fewer than this many extractable characters is treated as scanned
# and routed through OCR instead.
MIN_TEXT_CHARS_PER_PAGE = 40
OCR_DPI = 200

try:
    _enc = tiktoken.get_encoding("o200k_base")
except Exception:  # pragma: no cover
    _enc = tiktoken.get_encoding("cl100k_base")


def _count_tokens(text: str) -> int:
    return len(_enc.encode(text))


# ---------- extraction ----------
def _ocr_image(img: Image.Image) -> str:
    return pytesseract.image_to_string(img, lang="fas")


def _extract_pdf(path: str):
    """Yield (page_number, text). Falls back to OCR for scanned pages."""
    used_ocr = False
    doc = fitz.open(path)
    for i, page in enumerate(doc, start=1):
        text = page.get_text("text").strip()
        if len(text) < MIN_TEXT_CHARS_PER_PAGE:
            pix = page.get_pixmap(dpi=OCR_DPI)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            text = _ocr_image(img).strip()
            if text:
                used_ocr = True
        if text:
            yield i, text
    doc.close()
    return used_ocr


def _is_ruling_style(style_name: str) -> bool:
    """True for the paragraph style that marks a numbered ruling (مسأله).
    Matched loosely because authors name the style differently."""
    n = normalize_fa(style_name or "")
    return "مساله" in n or "مسیله" in n


def _heading_level(style_name: str):
    """Return the heading depth for 'Heading N' / 'عنوان N' styles, else None."""
    n = (style_name or "").strip()
    m = re.match(r"^(?:Heading|heading|عنوان)\s*([1-9])$", n)
    return int(m.group(1)) if m else None


def _docx_tables_text(d):
    out = []
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                out.append(" | ".join(cells))
    return out


def _extract_docx(path: str):
    """Structure-aware DOCX extraction.

    Religious/legal books (رساله) mark each ruling with a dedicated paragraph
    style and continue it in following body paragraphs, while the ruling NUMBER
    exists only implicitly (by order). Reading paragraphs as one flat blob
    therefore separates a ruling from its own clauses and loses its number.

    Here each ruling becomes ONE self-contained block:
        مسأله 1168 (احکام نماز ← شكّيات ← شکّ‌های مبطل)
        شکّ‌هايى كه نماز را باطل مى‏كند از اين قرار است:
        اوّل: ...  دوم: شکّ در شماره ركعتهاى نماز سه ركعتى. ...

    Documents without that style fall back to grouping by heading (useful for
    Q&A files), and documents without headings yield a single blob as before.
    """
    d = DocxDocument(path)
    paras = [p for p in d.paragraphs if p.text and p.text.strip()]
    has_rulings = any(_is_ruling_style(p.style.name) for p in paras)
    has_headings = any(_heading_level(p.style.name) for p in paras)

    if not has_rulings and not has_headings:
        parts = [p.text.strip() for p in paras] + _docx_tables_text(d)
        text = "\n".join(parts).strip()
        if text:
            yield None, text
        return

    heads = {}          # heading level -> latest text
    blocks = []         # list of (title, [lines])
    cur = None
    counter = 0

    def breadcrumb():
        return " ← ".join(heads[k] for k in sorted(heads) if heads.get(k))

    for p in paras:
        text = p.text.strip()
        lvl = _heading_level(p.style.name)
        if lvl:
            heads[lvl] = text
            for deeper in [k for k in heads if k > lvl]:
                heads.pop(deeper, None)
            cur = None                      # a heading ends the previous block
            continue
        if _is_ruling_style(p.style.name):
            counter += 1
            crumb = breadcrumb()
            title = f"مسأله {counter}" + (f" ({crumb})" if crumb else "")
            cur = (title, [text])
            blocks.append(cur)
        elif cur is not None:
            cur[1].append(text)             # continuation of the current ruling
        else:
            crumb = breadcrumb()
            cur = (crumb, [text]) if crumb else ("", [text])
            blocks.append(cur)

    for title, lines in blocks:
        body = "\n".join(lines).strip()
        if not body:
            continue
        yield None, (f"{title}\n{body}" if title else body)

    tables = _docx_tables_text(d)
    if tables:
        yield None, "\n".join(tables)


def _extract_image(path: str):
    text = _ocr_image(Image.open(path)).strip()
    if text:
        yield None, text


def extract(path: str, mime: str):
    """Return (list_of_(page,text), used_ocr)."""
    used_ocr = False
    pages = []
    if mime == "application/pdf" or path.lower().endswith(".pdf"):
        gen = _extract_pdf(path)
        try:
            while True:
                pages.append(next(gen))
        except StopIteration as stop:
            used_ocr = bool(stop.value)
    elif path.lower().endswith((".docx",)):
        pages = list(_extract_docx(path))
    elif path.lower().endswith((".png", ".jpg", ".jpeg", ".webp", ".tiff", ".bmp")):
        pages = list(_extract_image(path))
        used_ocr = True
    else:
        raise ValueError(f"نوع فایل پشتیبانی نمی‌شود: {mime}")
    return pages, used_ocr


# ---------- chunking ----------
# Split just before a structural heading like "مساله 1168" / "ماده 5", or before
# a Q&A item ("سوال 3:" / "پرسش:") so each question keeps its own answer. Text is
# already normalized, so 'مسأله' has become 'مساله'. Anchored to line starts so an
# in-text cross-reference ("تفصيل اين مسأله در مسأله 1202 مى‏آيد") never cuts a
# ruling in half; the digit/colon lookaheads skip prose like "در این مساله".
_STRUCT_SPLIT = re.compile(
    r"(?m)^(?="
    r"(?:مساله|ماده)\s*[0-9]{1,5}"
    r"|(?:سوال|پرسش)\s*[0-9]{0,4}\s*[:：]"
    r")"
)


def _window(tokens, page, max_tokens, step, out):
    start = 0
    while start < len(tokens):
        window = tokens[start:start + max_tokens]
        s = _enc.decode(window).strip()
        if s:
            out.append((s, page, len(window)))
        start += step


def chunk_text(text: str, page, max_tokens: int, overlap: int):
    """Structure-aware chunking: split on 'مساله N' / 'ماده N' headings so each
    ruling stays whole (heading + all its clauses in one chunk), then window any
    oversized segment. Falls back to plain token windows for text that has no
    such headings (intros, tables of contents, etc.)."""
    step = max(1, max_tokens - overlap)
    segments = [s for s in _STRUCT_SPLIT.split(text) if s.strip()]
    if len(segments) <= 1:
        segments = [text]
    chunks = []
    for seg in segments:
        seg = seg.strip()
        if not seg:
            continue
        toks = _enc.encode(seg)
        if len(toks) <= max_tokens:
            chunks.append((seg, page, len(toks)))
        else:
            _window(toks, page, max_tokens, step, chunks)
    return chunks


# ---------- main task ----------
def ingest_document(document_id: int):
    """Called by the worker. Extracts, chunks, embeds and stores the document."""
    db = SessionLocal()
    try:
        doc = db.query(Document).get(document_id)
        if not doc:
            return
        doc.status = "processing"
        db.commit()

        max_tokens = int(get_setting(db, "chunk_tokens", 800))
        overlap = int(get_setting(db, "chunk_overlap", 150))
        emb_model = get_setting(db, "embedding_model", "text-embedding-3-large")
        emb_dims = int(get_setting(db, "embedding_dims", 1536))

        pages, used_ocr = extract(doc.stored_path, doc.mime)
        # canonicalize to Persian forms so stored chunks + embeddings match
        # Persian-typed queries (ك→ک, ي/ى→ی, ة→ه, strip harakat, digits)
        pages = [(pg, normalize_fa(txt)) for pg, txt in pages]

        all_chunks = []
        for page_no, text in pages:
            all_chunks.extend(chunk_text(text, page_no, max_tokens, overlap))

        if not all_chunks:
            doc.status = "error"
            doc.error = "هیچ متنی از فایل استخراج نشد."
            db.commit()
            return

        # embed in batches
        BATCH = 64
        for i in range(0, len(all_chunks), BATCH):
            batch = all_chunks[i:i + BATCH]
            vectors, _ = embed_texts([c[0] for c in batch], emb_model, emb_dims)
            for (content, page_no, tok), vec in zip(batch, vectors):
                db.add(Chunk(
                    document_id=doc.id,
                    content=content,
                    page=page_no,
                    token_count=tok,
                    embedding=vec,
                ))
            db.commit()

        doc.pages = len({p for p, _ in pages if p is not None})
        doc.chunk_count = len(all_chunks)
        doc.used_ocr = used_ocr
        doc.status = "ready"
        doc.error = None
        db.commit()
        log.info("Document %s ingested: %s chunks", doc.id, len(all_chunks))
    except Exception as e:  # pragma: no cover
        log.exception("ingest failed")
        doc = db.query(Document).get(document_id)
        if doc:
            doc.status = "error"
            doc.error = str(e)[:1000]
            db.commit()
    finally:
        db.close()
